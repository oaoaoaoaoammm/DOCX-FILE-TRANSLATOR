from docx import Document
from deep_translator import GoogleTranslator
import os

# Инициализация переводчика
translator = GoogleTranslator(source='ru', target='zh-CN')


def translate_text(text):
    return translator.translate(text)


def translate_docx(input_file):
    # Открываем документ
    doc = Document(input_file)

    # Переводим текст в абзацах
    for para in doc.paragraphs:
        if para.text.strip():
            para.text = translate_text(para.text)

    # Переводим текст внутри таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    cell.text = translate_text(cell.text)

    # Получаем название файла без расширения
    base_name = os.path.splitext(os.path.basename(input_file))[0]

    # Переводим название файла
    translated_name = translate_text(base_name)

    # Формируем новое название файла с добавлением суффикса "_translated"
    output_file = f"{translated_name}_translated.docx"

    # Сохраняем переведенный документ с новым именем
    doc.save(output_file)
    return output_file


# Пример использования
input_file = 'двери.docx'
output_file = translate_docx(input_file)
print(f"File translated and saved as: {output_file}")